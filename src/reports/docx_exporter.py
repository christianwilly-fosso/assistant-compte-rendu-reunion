"""Word (.docx) export of a meeting report."""

from __future__ import annotations

import logging
import tempfile
import uuid
from pathlib import Path

from src.reports.report_model import ActionItem, MeetingReport, UNSPECIFIED

_LOGGER = logging.getLogger(__name__)


class DocxExporter:
    """Persist a :class:`MeetingReport` as a Word document.

    Returns ``None`` and logs a warning if ``python-docx`` is unavailable,
    so that the rest of the pipeline keeps working.
    """

    def export(self, report: MeetingReport, suggested_name: str = "compte_rendu") -> str | None:
        try:
            from docx import Document
        except ImportError:
            _LOGGER.warning("python-docx n'est pas installé : export DOCX désactivé.")
            return None

        document = Document()
        document.add_heading("Compte rendu de réunion", level=1)

        document.add_heading("Informations générales", level=2)
        document.add_paragraph(f"Date : {report.date or UNSPECIFIED}")
        document.add_paragraph(
            "Participants : " + (", ".join(report.participants) if report.participants else UNSPECIFIED)
        )
        document.add_paragraph(f"Sujet : {report.subject or UNSPECIFIED}")

        document.add_heading("Résumé exécutif", level=2)
        document.add_paragraph(report.executive_summary or UNSPECIFIED)

        _add_numbered_section(document, "Points discutés", report.discussed_points)
        _add_numbered_section(document, "Décisions prises", report.decisions)

        document.add_heading("Actions à faire", level=2)
        _add_actions_table(document, report.action_items)

        _add_numbered_section(document, "Points en suspens", report.pending_points)

        document.add_heading("Transcription complète", level=2)
        document.add_paragraph(report.full_transcription or UNSPECIFIED)

        path = Path(tempfile.gettempdir()) / f"{suggested_name}_{uuid.uuid4().hex}.docx"
        document.save(str(path))
        return str(path)


def _add_numbered_section(document, title: str, items: list[str]) -> None:
    document.add_heading(title, level=2)
    cleaned = [item.strip() for item in items if item.strip()]
    if not cleaned:
        document.add_paragraph(UNSPECIFIED)
        return
    for index, item in enumerate(cleaned, start=1):
        document.add_paragraph(f"{index}. {item}")


def _add_actions_table(document, actions: list[ActionItem]) -> None:
    table = document.add_table(rows=1, cols=4)
    table.style = "Light List"
    header = table.rows[0].cells
    header[0].text = "Action"
    header[1].text = "Responsable"
    header[2].text = "Échéance"
    header[3].text = "Statut"

    if not actions:
        row = table.add_row().cells
        row[0].text = UNSPECIFIED
        row[1].text = UNSPECIFIED
        row[2].text = UNSPECIFIED
        row[3].text = UNSPECIFIED
        return

    for action in actions:
        row = table.add_row().cells
        row[0].text = action.description or UNSPECIFIED
        row[1].text = action.responsible or UNSPECIFIED
        row[2].text = action.due_date or UNSPECIFIED
        row[3].text = action.status or "À faire"
